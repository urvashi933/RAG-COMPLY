from pypdf import PdfReader
import docx
import os

# def extract_text(file_path):
#     ext = os.path.splitext(file_path)[1].lower()
#     text = ""
#     if ext == ".pdf":
#         reader = PdfReader(file_path)
#         for page in reader.pages:
#             text += page.extract_text() + "\n"
#     elif ext == ".docx":
#         doc = docx.Document(file_path)
#         for p in doc.paragraphs:
#             text += p.text + "\n"
#     elif ext == ".txt":
#         with open(file_path, 'r', encoding='utf-8') as f:
#             text = f.read()
#     return text

def load_document(file_path):
    text=""

    if file_path.endswith(".pdf"):
        reader = PdfReader(file_path)
        for page_no,page in enumerate(reader.pages):
            page_text=page.extract_text()
            if page_text:
                text+= f"\n [PAGE {page_no+1}]\n{page_text}"
    elif file_path.endswith(".docx"):
        doc=docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text+ "\n"

    elif file_path.endswith(".txt"):
        with open(file_path,"r",encoding="utf-8") as f:
            text = f.read()

    return text